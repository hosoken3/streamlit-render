# ======================================================
# app.py : 環境変数優先（Render）＋ secrets.toml フォールバック（ローカル）
#   - 認証: USERNAME / PASSWORD（なければ secrets から）
#   - Gemini: GEMINI_API_KEY（なければ secrets["api"]["gemini_key"]）
#   - 3タブ: ①マッチング実行 ②アイデア生成（Gemini）③ファイル作成
# ======================================================

import os
import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from PyPDF2 import PdfReader

# Gemini（新SDK）
# pip install google-genai
from google import genai

# ------------------------------------------------------
# 0. ページ設定
# ------------------------------------------------------
st.set_page_config(page_title="技術ニーズマッチング（Render環境変数対応版）", layout="wide")

# ------------------------------------------------------
# 1. シークレットの取得（環境変数を優先、無ければ secrets.toml）
# ------------------------------------------------------
def get_secret(key: str, default: str | None = None):
    """環境変数 → st.secrets の順で値を取得"""
    env_val = os.getenv(key)
    if env_val is not None and env_val != "":
        return env_val
    # st.secrets に階層がある場合は別途明示で扱う（下で実装）
    return default

# 認証情報
USERNAME = get_secret("USERNAME")
PASSWORD = get_secret("PASSWORD")

# フォールバック: secrets.toml（.streamlit/secrets.toml）
# 例:
# [auth]
# users = [
#   { username = "tanaka", password = "pass123" },
#   { username = "sato",   password = "pass456" }
# ]
if not USERNAME or not PASSWORD:
    # 複数ユーザー方式（配列）に対応。単一キー方式にも対応。
    try:
        auth_block = st.secrets.get("auth", {})
        # ① 単一キー（USERNAME / PASSWORD）での運用
        if not USERNAME:
            USERNAME = auth_block.get("username", USERNAME)
        if not PASSWORD:
            PASSWORD = auth_block.get("password", PASSWORD)

        # ② 複数ユーザー（users 配列）を許容：この場合は複数ユーザー認証に切り替え
        USERS_LIST = auth_block.get("users", None)  # [{username, password}, ...]
    except Exception:
        USERS_LIST = None
else:
    USERS_LIST = None  # 環境変数で単一ユーザー運用の場合は配列は使わない

# Gemini API キー（環境変数優先 → secrets["api"]["gemini_key"]）
GEMINI_API_KEY = get_secret("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    try:
        GEMINI_API_KEY = st.secrets.get("api", {}).get("gemini_key", "")
    except Exception:
        GEMINI_API_KEY = ""

# ------------------------------------------------------
# 2. ログイン（環境変数の単一ユーザー, または secrets の複数ユーザー）
# ------------------------------------------------------
def login_single_user():
    """USERNAME/PASSWORD での単一ユーザー認証"""
    st.title("🔒 ログイン")
    user = st.text_input("ユーザー名を入力してください")
    pw = st.text_input("パスワードを入力してください", type="password")
    if st.button("ログイン"):
        if user == USERNAME and pw == PASSWORD:
            st.session_state["logged_in"] = True
            st.session_state["user_name"] = user
            st.success("ログイン成功！")
            st.rerun()
        else:
            st.error("ユーザー名またはパスワードが違います。")

def login_multi_users(users_list: list[dict]):
    """secrets.toml の [auth].users を使った複数ユーザー認証"""
    st.title("🔒 ログイン")
    user = st.text_input("ユーザー名を入力してください")
    pw = st.text_input("パスワードを入力してください", type="password")
    if st.button("ログイン"):
        for u in users_list:
            if user == u.get("username") and pw == u.get("password"):
                st.session_state["logged_in"] = True
                st.session_state["user_name"] = user
                st.success(f"ようこそ、{user} さん！")
                st.rerun()
                return
        st.error("ユーザー名またはパスワードが違います。")

# ログイン状態の初期化
if "logged_in" not in st.session_state:
    st.session_state["logged_in"] = False

# ログインフロー
if not st.session_state["logged_in"]:
    # 環境変数に USERNAME/PASSWORD がある → 単一ユーザー認証
    # 無い場合、secrets の users 配列があれば複数ユーザー認証
    if USERNAME and PASSWORD:
        login_single_user()
    elif USERS_LIST:
        login_multi_users(USERS_LIST)
    else:
        st.error("認証情報が見つかりません。Render 環境変数（USERNAME/PASSWORD）または secrets.toml を設定してください。")
    st.stop()

# ログイン後の表示
st.sidebar.success(f"👤 ログイン中：{st.session_state['user_name']}")
if st.sidebar.button("ログアウト"):
    st.session_state["logged_in"] = False
    st.rerun()

# ------------------------------------------------------
# 3. CSV / PDF 読み込み
# ------------------------------------------------------
@st.cache_data
def load_csv(path="data/sample.csv"):
    try:
        return pd.read_csv(path)
    except Exception as e:
        st.warning(f"CSVの読み込みに失敗しました: {e}")
        return pd.DataFrame()

df = load_csv()

st.sidebar.header("📁 ファイル読込（任意）")
uploaded_csv = st.sidebar.file_uploader("CSVをアップロード", type=["csv"])
uploaded_pdf = st.sidebar.file_uploader("PDFをアップロード", type=["pdf"])

if uploaded_csv:
    try:
        df = pd.read_csv(uploaded_csv)
        st.sidebar.success("CSVを読み込みました。")
    except Exception as e:
        st.sidebar.error(f"CSV読込エラー: {e}")

pdf_text = ""
if uploaded_pdf:
    try:
        reader = PdfReader(uploaded_pdf)
        for page in reader.pages:
            pdf_text += page.extract_text() or ""
        st.sidebar.success("PDFを読み込みました。")
    except Exception as e:
        st.sidebar.error(f"PDF読込エラー: {e}")

# ------------------------------------------------------
# 4. メイン画面（3タブ構成）
# ------------------------------------------------------
tab1, tab2, tab3 = st.tabs(["①マッチング実行", "②アイデア生成（Gemini）", "③ファイル作成"])

# ---------------------------
# ① マッチング実行
# ---------------------------
with tab1:
    st.header("マッチング実行")
    col1, col2, col3 = st.columns([2, 2, 2])
    with col1:
        company = st.text_input("企業名：")
        news_kw = st.text_input("ニュース名：", placeholder="（曖昧検索）")
    with col2:
        major = st.selectbox("大分類：", ["", "材料", "機械", "電気"])
    with col3:
        middle = st.selectbox("中分類：", ["", "加工", "AI", "制御"])

    c1, c2 = st.columns([1, 1])
    with c1:
        if st.button("検索"):
            st.success("ダミー検索を実行しました。")
    with c2:
        if st.button("クリア"):
            st.rerun()

    st.markdown("### 結果一覧")
    if not df.empty:
        df_show = df.copy()
        if "番号" not in df_show.columns:
            df_show.insert(0, "番号", range(1, len(df_show) + 1))
        df_show.insert(1, "✔選択", False)
        st.dataframe(df_show, use_container_width=True, height=300)
    else:
        st.info("データがありません。左サイドバーからCSVをアップロードできます。")

# ---------------------------
# ② アイデア生成（Gemini）
# ---------------------------
with tab2:
    st.header("💡 Gemini によるアイデア生成")

    # キーチェック
    if not GEMINI_API_KEY:
        st.error("Gemini APIキーが設定されていません。Render 環境変数 GEMINI_API_KEY か secrets.toml の [api].gemini_key を設定してください。")
    else:
        # クライアント初期化
        client = genai.Client(api_key=GEMINI_API_KEY)

        st.write("以下のCSVデータまたはPDF内容をもとにAIが新しいアイデアを生成します。")
        if not df.empty:
            st.dataframe(df.head(5), use_container_width=True)
        elif pdf_text:
            st.info("PDFのテキストが読み込まれています。")
        else:
            st.info("CSVまたはPDFをアップロードしてください。")

        prompt = st.text_area(
            "🔧 プロンプト（AIへの指示文）",
            "以下の技術ニュースをもとに、新しい応用技術アイデアを3つ提案してください。"
        )

        if st.button("🚀 Geminiでアイデア生成"):
            with st.spinner("Geminiが考え中..."):
                try:
                    text_summary = ""
                    if not df.empty:
                        # 先頭3行の要約
                        text_summary = "\n".join(
                            df.head(3).astype(str).fillna("").apply(lambda row: " ".join(row), axis=1)
                        )
                    elif pdf_text:
                        text_summary = pdf_text[:1500]

                    full_prompt = f"{prompt}\n\n元データ:\n{text_summary}"

                    # モデル名は適宜更新可（例: "gemini-1.5-flash", "gemini-2.0-flash"）
                    resp = client.models.generate_content(
                        model="gemini-1.5-flash",
                        contents=full_prompt,
                    )
                    out = getattr(resp, "text", None) or getattr(resp, "output_text", "")
                    if not out:
                        out = str(resp)
                    st.success("💡 アイデア生成完了！")
                    st.write(out)
                except Exception as e:
                    st.error(f"Geminiエラー: {e}")

# ---------------------------
# ③ ファイル作成
# ---------------------------
with tab3:
    st.header("📄 Wordファイル作成（ダミー）")
    st.write("選択データとPDF内容からWordレポートを生成します。")

    def make_docx(df_in: pd.DataFrame, pdf_text_in: str) -> bytes:
        doc = Document()
        doc.add_heading("技術ニーズ マッチング レポート（ダミー）", level=1)
        doc.add_paragraph(f"■ ログインユーザー：{st.session_state['user_name']}")
        doc.add_paragraph("■ PDF抽出テキスト（先頭100～150文字）：")
        doc.add_paragraph((pdf_text_in or "（PDF未読込）")[:150])

        if not df_in.empty:
            doc.add_heading("■ データ要約", level=2)
            for _, row in df_in.head(10).iterrows():
                title = str(row.get("技術ニュース名", row.get("タイトル", "（無題）")))
                company = str(row.get("企業名", ""))
                summary = str(row.get("要約", ""))
                doc.add_paragraph(f"・{title} / {company}")
                if summary:
                    doc.add_paragraph(f"  - 要約: {summary}")
        else:
            doc.add_paragraph("データがありません。")

        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.read()

    if st.button("Word出力"):
        content = make_docx(df, pdf_text)
        st.download_button(
            "📄 ダウンロード（output.docx）",
            data=content,
            file_name="output.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
